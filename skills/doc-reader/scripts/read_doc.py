"""
Document Reader — extract text/data from PDF, DOC, DOCX, TXT, RTF, XLS, XLSX.
Usage: python read_doc.py <filepath> [--json] [--sheet N] [--page N]

Auto-detects format from file extension and available libraries.
"""
import sys
import os
import json
import re
import subprocess
from pathlib import Path


def read_txt(path):
    """Read plain text / RTF — strip RTF markup if needed."""
    raw = path.read_text(encoding="utf-8", errors="replace")
    if path.suffix.lower() == ".rtf":
        return _strip_rtf(raw)
    return raw


def _strip_rtf(rtf):
    """Basic RTF to plain text."""
    text = re.sub(r'\\[a-z]+\d*\s?', '', rtf)
    text = re.sub(r'[{}]', '', text)
    text = text.replace('\\par', '\n')
    text = text.replace('\\tab', '\t')
    text = re.sub(r'\\\'[0-9a-fA-F]{2}', '', text)
    return text.strip()


def read_docx(path):
    """Read .docx using python-docx."""
    import docx
    doc = docx.Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    # Tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        parts.append("\n--- Table ---\n" + "\n".join(rows))
    return "\n".join(parts)


def read_doc(path):
    """Read old .doc format — try catdoc, then binary fallback."""
    # Try catdoc
    try:
        r = subprocess.run(["catdoc", str(path)], capture_output=True, text=True, timeout=10)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    # Try antiword
    try:
        r = subprocess.run(["antiword", str(path)], capture_output=True, text=True, timeout=10)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    # Binary fallback — extract readable text
    raw = path.read_bytes()
    # Try to find text between null bytes
    text_parts = []
    current = []
    for b in raw:
        if 32 <= b < 127 or b in (10, 13, 9):
            current.append(chr(b))
        else:
            if len(current) > 3:
                text_parts.append("".join(current))
            current = []
    if current:
        text_parts.append("".join(current))
    result = "\n".join(text_parts)
    if len(result) < 50:
        return "[Could not extract meaningful text from .doc file. Try installing catdoc or antiword.]"
    return result


def read_pdf(path):
    """Read PDF — try pdftotext, then PyPDF2, then pdfplumber."""
    # Try pdftotext
    try:
        r = subprocess.run(["pdftotext", str(path), "-"], capture_output=True, text=True, timeout=30)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    # Try PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(str(path))
        parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                parts.append(f"[Page {i+1}]\n{text}")
        if parts:
            return "\n\n".join(parts)
    except ImportError:
        pass
    # Try pdfplumber
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    parts.append(f"[Page {i+1}]\n{text}")
        if parts:
            return "\n\n".join(parts)
    except ImportError:
        pass
    return "[Cannot read PDF. Install pdftotext (poppler-utils) or PyPDF2/pdfplumber.]"


def read_xlsx(path, sheet=None):
    """Read .xlsx using openpyxl."""
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    sheets = [wb.active] if sheet is None else [wb.worksheets[sheet]]
    parts = []
    for ws in sheets:
        parts.append(f"[Sheet: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(cell) if cell is not None else "" for cell in row]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def read_xls(path, sheet=None):
    """Read old .xls using xlrd."""
    import xlrd
    wb = xlrd.open_workbook(str(path))
    sheets = [wb.sheet_by_index(0)] if sheet is None else [wb.sheet_by_index(sheet)]
    parts = []
    for ws in sheets:
        parts.append(f"[Sheet: {ws.name}]")
        for row_idx in range(ws.nrows):
            cells = [str(ws.cell_value(row_idx, col)) for col in range(ws.ncols)]
            if any(c.strip() for c in cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def main():
    if len(sys.argv) < 2:
        print("Usage: python read_doc.py <filepath> [--json] [--sheet N] [--page N]")
        sys.exit(1)

    filepath = Path(sys.argv[1])
    if not filepath.exists():
        print(f"File not found: {filepath}")
        sys.exit(1)

    as_json = "--json" in sys.argv
    sheet_idx = None
    for i, arg in enumerate(sys.argv):
        if arg == "--sheet" and i + 1 < len(sys.argv):
            sheet_idx = int(sys.argv[i + 1])

    ext = filepath.suffix.lower()
    result = ""

    try:
        if ext == ".pdf":
            result = read_pdf(filepath)
        elif ext == ".docx":
            result = read_docx(filepath)
        elif ext == ".doc":
            result = read_doc(filepath)
        elif ext in (".txt", ".rtf", ".md", ".csv", ".log", ".py", ".js", ".html", ".css", ".json", ".xml", ".yaml", ".yml"):
            result = read_txt(filepath)
        elif ext == ".xlsx":
            result = read_xlsx(filepath, sheet_idx)
        elif ext == ".xls":
            result = read_xls(filepath, sheet_idx)
        else:
            # Try as text
            result = read_txt(filepath)
    except Exception as e:
        result = f"Error reading {filepath}: {e}"

    if as_json:
        print(json.dumps({"file": str(filepath), "format": ext, "content": result[:50000]},
                         ensure_ascii=False, indent=2))
    else:
        print(result[:50000])


if __name__ == "__main__":
    main()
